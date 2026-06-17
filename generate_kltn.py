#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate KLTN Word document from Mau 1.docx template."""

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "/Users/tranvietduc/Downloads/Mau 1.docx"
OUTPUT = "/Users/tranvietduc/Downloads/KLTN_Tran_Viet_Duc_P2P_Chat.docx"

STUDENT = "TRẦN VIỆT ĐỨC"
STUDENT_TITLE = "Trần Việt Đức"
MSV = "24260029"
LOP = "K6TXCNTT01"
ADVISOR = "ThS. Đỗ Thành Công"
TOPIC = (
    "NGHIÊN CỨU VÀ XÂY DỰNG ỨNG DỤNG NHẮN TIN BẢO MẬT "
    "THỜI GIAN THỰC DỰA TRÊN KIẾN TRÚC PEER-TO-PEER VỚI MÃ HÓA ĐẦU CUỐI"
)
DATE_SIGN = "Hà Nội, ngày 17 tháng 6 năm 2026"
FONT = "Times New Roman"


def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_run(run, bold=None, size=None, name=FONT):
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = size
    if name:
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
        rfonts.set(qn("w:eastAsia"), name)


def add_run(paragraph, text, bold=None, size=None):
    run = paragraph.add_run(text)
    set_run(run, bold=bold, size=size)
    return run


def add_p(doc, text="", style="Normal", align=None, first_indent=None,
          space_before=None, space_after=None, line_spacing=None, bold=False):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run(run, bold=bold, size=Pt(13))
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
    for r in p.runs:
        if r.font.size is None:
            set_run(r, size=Pt(13))
    return p


def add_body(doc, text, first_indent=Emu(360045)):
    return add_p(doc, text, first_indent=first_indent, line_spacing=1.5)


def add_heading(doc, text, level=1):
    sizes = {1: Pt(16), 2: Pt(14), 3: Pt(13)}
    p = add_p(doc, text, bold=True, space_before=Pt(12), space_after=Pt(6))
    for r in p.runs:
        set_run(r, bold=True, size=sizes.get(level, Pt(13)))
    return p


def add_bullets(doc, items):
    for item in items:
        add_p(doc, f"• {item}", style="List Paragraph", line_spacing=1.5)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_toc_tab_stop(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.insert(0, tabs)
    else:
        for t in list(tabs.findall(qn("w:tab"))):
            tabs.remove(t)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9014")
    tabs.append(tab)


def add_toc_lines(doc, lines):
    text = "\n".join(f"{title}\t{page}" for title, page in lines)
    p = add_p(doc, text, space_before=Pt(6), space_after=Pt(6))
    set_toc_tab_stop(p)
    for r in p.runs:
        set_run(r, size=Pt(13))


def add_glossary_table(doc):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Thuật ngữ"
    hdr[1].text = "Giải thích"
    terms = [
        ("E2E", "End-to-End Encryption – Mã hóa đầu cuối"),
        ("P2P", "Peer-to-Peer – Kiến trúc ngang hàng"),
        ("WebRTC", "Web Real-Time Communication"),
        ("STUN", "Session Traversal Utilities for NAT"),
        ("ICE", "Interactive Connectivity Establishment"),
        ("SDP", "Session Description Protocol"),
        ("AES-GCM", "Advanced Encryption Standard – Galois/Counter Mode"),
        ("X25519", "Thuật toán trao đổi khóa ECDH trên đường cong elliptic"),
        ("JWT", "JSON Web Token – Xác thực phiên làm việc"),
        ("Signaling", "Kênh trao đổi metadata thiết lập kết nối P2P"),
    ]
    for term, definition in terms:
        row = tbl.add_row().cells
        row[0].text = term
        row[1].text = definition


def add_cover(doc):
    for text in [
        "BỘ GIÁO DỤC VÀ ĐÀO TẠO",
        "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ VÀ QUẢN LÝ HỮU NGHỊ",
        "KHOA CÔNG NGHỆ THÔNG TIN",
    ]:
        p = add_p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)
        for r in p.runs:
            set_run(r, bold=True, size=Pt(14))
    for _ in range(4):
        add_p(doc, "")
    p = add_p(doc, "KHÓA LUẬN TỐT NGHIỆP", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_run(r, bold=True, size=Pt(18))
    add_p(doc, "")
    p = add_p(doc, f"TÊN ĐỀ TÀI: {TOPIC}", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_run(r, bold=True, size=Pt(13))
    for _ in range(3):
        add_p(doc, "")
    for line in [
        f"Cán bộ hướng dẫn: {ADVISOR}",
        f"Sinh viên: {STUDENT}",
        f"Mã sinh viên: {MSV}",
        "Chuyên ngành: Công nghệ thông tin",
        f"Lớp: {LOP}",
    ]:
        p = add_p(doc, line)
        for r in p.runs:
            set_run(r, bold=True, size=Pt(13))
    for _ in range(4):
        add_p(doc, "")
    p = add_p(doc, "Hà Nội - 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_run(r, bold=True, size=Pt(13))


def add_loi_cam_on(doc):
    add_heading(doc, "LỜI CẢM ƠN")
    paras = [
        f"Em xin gửi lời cảm ơn chân thành tới các thầy cô Trường Đại học Công nghệ và Quản lý Hữu Nghị, "
        f"đặc biệt là thầy {ADVISOR} đã tận tình hướng dẫn em trong suốt quá trình thực hiện khóa luận tốt nghiệp.",
        "Em cũng xin cảm ơn các thầy cô Khoa Công nghệ thông tin đã truyền đạt kiến thức nền tảng về mạng máy tính, "
        "an toàn thông tin và lập trình di động, tạo tiền đề quan trọng để em hoàn thành đề tài.",
        "Do hạn chế về thời gian và kinh nghiệm, khóa luận không tránh khỏi thiếu sót. "
        "Em rất mong nhận được ý kiến đóng góp của thầy cô để hoàn thiện hơn.",
    ]
    for t in paras:
        add_body(doc, t)
    add_p(doc, "Em xin chân thành cảm ơn!", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Sinh viên thực hiện", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "")
    p = add_p(doc, STUDENT_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_run(r, bold=True, size=Pt(14))
    add_page_break(doc)


def add_muc_luc(doc):
    add_heading(doc, "MỤC LỤC", level=1)
    add_toc_lines(doc, [
        ("LỜI CẢM ƠN", "2"),
        ("THÔNG TIN KẾT QUẢ ĐỒ ÁN", "4"),
        ("DANH SÁCH HÌNH ẢNH", "6"),
        ("GIẢI THÍCH THUẬT NGỮ", "7"),
        ("MỞ ĐẦU", "8"),
        ("CHƯƠNG 1. CƠ SỞ LÝ THUYẾT", "14"),
        ("CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "28"),
        ("CHƯƠNG 3. CÀI ĐẶT, TRIỂN KHAI VÀ THỬ NGHIỆM", "42"),
        ("KẾT LUẬN VÀ KIẾN NGHỊ", "56"),
        ("TÀI LIỆU THAM KHẢO", "58"),
    ])
    add_page_break(doc)


def add_thong_tin_do_an(doc):
    add_heading(doc, "THÔNG TIN KẾT QUẢ ĐỒ ÁN")
    add_heading(doc, "Thông tin chung", level=2)
    add_body(doc, f"Tên đề tài: {TOPIC.title()}")
    add_body(doc, f"Giảng viên hướng dẫn: {ADVISOR}")
    add_body(doc, f"Sinh viên thực hiện: {STUDENT_TITLE}")
    add_heading(doc, "Mục tiêu", level=2)
    add_bullets(doc, [
        "Nghiên cứu cơ chế mã hóa đầu cuối và kiến trúc Peer-to-Peer trong nhắn tin thời gian thực.",
        "Xây dựng ứng dụng Flutter kết nối P2P qua WebRTC DataChannel.",
        "Triển khai signaling server không lưu trữ nội dung tin nhắn.",
        "Đánh giá tính bảo mật, độ trễ và khả năng vận hành thực tế.",
    ])
    add_heading(doc, "Nội dung chính", level=2)
    add_bullets(doc, [
        "Tổng hợp lý thuyết E2E, P2P, WebRTC, STUN/ICE.",
        "Phân tích yêu cầu, thiết kế kiến trúc và giao thức signaling.",
        "Cài đặt module mã hóa X25519 + AES-GCM trên Flutter.",
        "Triển khai server Node.js và kiểm thử luồng chat P2P.",
    ])
    add_heading(doc, "Kết quả chính đạt được", level=2)
    add_bullets(doc, [
        "Hoàn thiện ứng dụng nhắn tin bảo mật trên Flutter.",
        "Server xác nhận messagesStored = 0 – không lưu tin nhắn tập trung.",
        "Tin nhắn được mã hóa trước khi truyền qua kênh P2P.",
        "Lịch sử chat lưu cục bộ trên thiết bị người dùng.",
    ])
    add_page_break(doc)


def add_mo_dau(doc):
    add_heading(doc, "MỞ ĐẦU")
    add_heading(doc, "1. Đặt vấn đề", level=2)
    add_body(doc,
        "Trong thời đại số hóa, ứng dụng nhắn tin trở thành kênh giao tiếp không thể thiếu. "
        "Tuy nhiên, hầu hết nền tảng phổ biến lưu trữ tin nhắn trên máy chủ trung tâm, khiến người dùng "
        "phụ thuộc vào nhà cung cấp dịch vụ và chịu rủi ro rò rỉ dữ liệu, truy cập trái phép hoặc phân tích hành vi.")
    add_body(doc,
        "Các sự cố bảo mật trong lĩnh vực truyền thông cho thấy mô hình client-server truyền thống "
        "khó đảm bảo quyền riêng tư tuyệt đối. Người dùng ngày càng quan tâm đến giải pháp "
        "mã hóa đầu cuối (End-to-End Encryption – E2E) và kiến trúc phân tán, nơi nội dung hội thoại "
        "chỉ có thể đọc được tại thiết bị của người gửi và người nhận.")
    add_body(doc,
        "Đề tài của em hướng tới xây dựng hệ thống nhắn tin thời gian thực ưu tiên quyền riêng tư: "
        "tin nhắn được mã hóa E2E bằng X25519 và AES-GCM, truyền trực tiếp qua WebRTC DataChannel theo mô hình P2P, "
        "trong khi server chỉ đảm nhiệm signaling (trao đổi SDP, ICE candidate, public key) mà không lưu nội dung tin nhắn.")
    add_heading(doc, "2. Mục đích nghiên cứu", level=2)
    add_bullets(doc, [
        "Nghiên cứu lý thuyết về mã hóa đầu cuối và kiến trúc Peer-to-Peer.",
        "Phân tích và thiết kế ứng dụng chat bảo mật với Flutter và signaling server Node.js.",
        "Cài đặt, triển khai thử nghiệm và đánh giá hiệu năng cùng mức độ bảo mật.",
    ])
    add_heading(doc, "3. Tình hình nghiên cứu", level=2)
    add_body(doc,
        "Signal, WhatsApp và Matrix là các hệ thống tiêu biểu áp dụng E2E. Signal Protocol sử dụng Double Ratchet "
        "kết hợp X3DH. Matrix kết hợp federation server với Olm/Megolm. Briar là ứng dụng P2P hoàn toàn qua Bluetooth/Wi-Fi/Tor, "
        "không phụ thuộc server trung tâm. WebRTC được chuẩn hóa bởi W3C/IETF, hỗ trợ truyền media và data channel realtime.")
    add_body(doc,
        "Trong giới học thuật và cộng đồng mã nguồn mở, xu hướng privacy-by-design và local-first storage "
        "ngày càng phổ biến. Tuy nhiên, tại Việt Nam, số lượng đề tài tốt nghiệp kết hợp đồng thời Flutter, WebRTC P2P "
        "và E2E với chính sách không lưu tin nhắn server còn hạn chế, tạo động lực cho em nghiên cứu đề tài này.")
    add_heading(doc, "4. Đối tượng, phạm vi và phương pháp", level=2)
    add_heading(doc, "4.1. Đối tượng nghiên cứu", level=3)
    add_bullets(doc, [
        "Cơ chế mã hóa đầu cuối trong nhắn tin.",
        "Kiến trúc P2P và giao thức WebRTC.",
        "Signaling server tối giản không lưu trữ tin nhắn.",
    ])
    add_heading(doc, "4.2. Phạm vi nghiên cứu", level=3)
    add_bullets(doc, [
        "Ứng dụng Flutter trên Android (có thể mở rộng iOS).",
        "Tin nhắn văn bản realtime, chưa hỗ trợ file đính kèm và nhóm chat.",
        "Signaling server Node.js + Socket.IO, STUN công khai của Google.",
    ])
    add_heading(doc, "4.3. Phương pháp nghiên cứu", level=3)
    add_bullets(doc, [
        "Nghiên cứu tài liệu: RFC, tài liệu WebRTC, cryptography package.",
        "Phân tích và thiết kế: UML use case, sequence diagram, kiến trúc phân tầng.",
        "Thực nghiệm: cài đặt hệ thống, kiểm thử chức năng và đánh giá bảo mật.",
    ])
    add_page_break(doc)


# Import chapter content from separate module to keep file manageable
from kltn_chapters import add_chuong_1, add_chuong_2, add_chuong_3, add_ket_luan


def main():
    doc = Document(TEMPLATE)
    clear_body(doc)

    add_cover(doc)
    add_page_break(doc)
    add_loi_cam_on(doc)
    add_muc_luc(doc)
    add_thong_tin_do_an(doc)
    add_heading(doc, "DANH SÁCH HÌNH ẢNH")
    add_bullets(doc, [
        "Hình 1.1. Kiến trúc tổng quan hệ thống P2P E2E",
        "Hình 2.1. Biểu đồ use case tổng quát",
        "Hình 2.2. Sequence diagram thiết lập kết nối P2P",
        "Hình 2.3. Sơ đồ kiến trúc phân tầng Flutter app",
        "Hình 2.4. Thiết kế CSDL local SQLite",
        "Hình 3.1. Giao diện đăng nhập/đăng ký",
        "Hình 3.2. Danh sách người dùng online",
        "Hình 3.3. Màn hình chat P2P với trạng thái E2E",
        "Hình 3.4. Kết quả kiểm thử API server (messagesStored = 0)",
    ])
    add_page_break(doc)
    add_heading(doc, "GIẢI THÍCH THUẬT NGỮ")
    add_glossary_table(doc)
    add_page_break(doc)
    add_mo_dau(doc)
    add_chuong_1(doc, add_heading, add_body, add_bullets, add_page_break)
    add_chuong_2(doc, add_heading, add_body, add_bullets, add_page_break)
    add_chuong_3(doc, add_heading, add_body, add_bullets, add_page_break)
    add_ket_luan(doc, add_heading, add_body, add_bullets, add_page_break)

    doc.save(OUTPUT)
    print(f"Đã tạo file: {OUTPUT}")


if __name__ == "__main__":
    main()
